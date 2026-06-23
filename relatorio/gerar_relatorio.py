"""
Gera o relatório final em DOCX (python-docx).
python relatorio/gerar_relatorio.py
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = Path(__file__).parent / "Relatorio_AML_FT_CloudWalk.docx"

# ── Paleta de cores ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x5C)   # cabeçalhos de tabela
BLUE   = RGBColor(0x2E, 0x75, 0xB6)   # headings H1
GRAY_L = RGBColor(0xD9, 0xE1, 0xF2)   # fundo linha par de tabela
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
RED    = RGBColor(0xC0, 0x00, 0x00)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, color=None, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    if color:
        run.font.color.rgb = color


def add_table(doc, headers, rows, col_widths_cm):
    """Cria tabela formatada com cabeçalho navy + listras alternadas."""
    n_cols = len(headers)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    hdr.height = Cm(0.75)
    for i, (h, w) in enumerate(zip(headers, col_widths_cm)):
        c = hdr.cells[i]
        c.width = Cm(w)
        set_cell_bg(c, "1A3A5C")
        cell_text(c, h, bold=True, color=WHITE, size=Pt(9.5), align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        bg = "D9E1F2" if ridx % 2 == 0 else "FFFFFF"
        for i, (val, w) in enumerate(zip(row_data, col_widths_cm)):
            c = row.cells[i]
            c.width = Cm(w)
            set_cell_bg(c, bg)
            bold = "**" in val
            text = val.replace("**", "")
            cell_text(c, text, bold=bold, color=BLACK, size=Pt(9.5))
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = "Normal"
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
        # underline via border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "2E75B6")
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK
        run.italic = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p


def add_body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p


def add_bullet(doc, text, bold=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_highlight_box(doc, text, color_hex="D9E1F2"):
    """Parágrafo com fundo colorido (blockquote/destaque)."""
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    run.bold = True
    return p


doc = Document()

# ── Configurar página A4 ──────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# ── Estilo de corpo padrão ────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)

# ── Footer: confidencial + nº página ─────────────────────────────────────
footer = section.footer
fp = footer.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = fp.add_run("Confidencial — Uso Interno — AML/PLD-FT  |  Página ")
run_f.font.size = Pt(9)
run_f.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
# adicionar número de página
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.text = "PAGE"
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")
run_pg = fp.add_run()
run_pg._r.append(fldChar1)
run_pg._r.append(instrText)
run_pg._r.append(fldChar2)
run_pg.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════════════════════
# PÁGINA DE ROSTO
# ═══════════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("RELATÓRIO FINAL")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = BLUE

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run("Sistema AML-FT CloudWalk")
r2.font.size = Pt(18)
r2.font.bold = True
r2.font.color.rgb = NAVY

doc.add_paragraph()

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_desc.add_run("Detecção de Lavagem de Dinheiro e Financiamento ao Terrorismo")
r3.font.size = Pt(13)
r3.font.italic = True
r3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

p_rail = doc.add_paragraph()
p_rail.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p_rail.add_run("PIX  ·  Cartão  ·  Wire")
r4.font.size = Pt(12)
r4.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

for _ in range(4):
    doc.add_paragraph()

# Caixa de metadados
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Autor",          "Victor OS — viictoros20@gmail.com"),
    ("Data",           "21 de junho de 2026"),
    ("Repositório",    "github.com/VictorOS1997/aml-case-cloudwalk"),
    ("Classificação",  "CONFIDENCIAL — Uso Interno — AML/PLD-FT"),
]
for i, (k, v) in enumerate(meta_data):
    c0 = meta_table.rows[i].cells[0]
    c1 = meta_table.rows[i].cells[1]
    c0.width = Cm(4)
    c1.width = Cm(12)
    set_cell_bg(c0, "1A3A5C")
    cell_text(c0, k, bold=True, color=WHITE, size=Pt(10.5))
    if k == "Classificação":
        set_cell_bg(c1, "FFF2CC")
        cell_text(c1, v, bold=True, color=RED, size=Pt(10.5))
    else:
        set_cell_bg(c1, "EEF3FB")
        cell_text(c1, v, size=Pt(10.5))

# Quebra de página
doc.add_paragraph().add_run().add_break(
    __import__("docx.enum.text", fromlist=["WD_BREAK_TYPE"]).WD_BREAK_TYPE.PAGE  # type: ignore
)

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 1 — RESUMO EXECUTIVO
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Resumo Executivo", 1)

add_body(doc, (
    "Este relatório descreve o sistema AML-FT desenvolvido para detectar, investigar e recomendar "
    "ações regulatórias sobre atividades suspeitas de lavagem de dinheiro e financiamento ao "
    "terrorismo na base de transações CloudWalk — 52.000 transações, R$ 230 milhões movimentados "
    "em PIX, Cartão e Wire entre julho e outubro de 2025."
))

add_heading(doc, "Resultados em síntese:", 3)
add_table(
    doc,
    headers=["Dimensão", "Resultado"],
    rows=[
        ["Base analisada",          "52.000 tx · 2.500 perfis KYC · 1.000 merchants · 3.310 senders ativos"],
        ["Regras implementadas",    "**22 regras** em 10 frentes de risco"],
        ["Alertas gerados",         "**10.576** (17 de 22 regras dispararam)"],
        ["Entidades no ranking",    "**4.156** (3.310 clientes + 846 merchants)"],
        ["Suspeitos detalhados",    "**Top 30** clientes (11 tipologias distintas)"],
        ["SAR elaborado",           "**1 SAR completo** — C101208 (score 22/22)"],
        ["Modelo ML (teste único)", "ROC-AUC **0,979** · PR-AUC **0,536** · Recall 93%"],
        ["Modelo ML (CV 5-fold)",   "ROC-AUC **0,956 ± 0,007** · PR-AUC **0,313 ± 0,048** (~9,6x baseline)"],
        ["Caso principal",          "C101208 — 7 tipologias · 11,5× renda · sanção confirmada"],
    ],
    col_widths_cm=[5.5, 10.5],
)

doc.add_paragraph()
add_body(doc, (
    "A abordagem combinou regras determinísticas (cobertura ampla), rótulo fraco (weak label das "
    "5 regras-core de alta confiança — R03, R07, R09, R12, R15) e XGBoost + Isolation Forest "
    "(generalização e detecção de anomalias não cobertas por regras). Um pipeline multi-agente "
    "LLM de 5 estágios orquestra o ciclo da detecção à decisão regulatória."
))

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 2 — ACHADOS SUSPEITOS E SAR
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Achados Suspeitos e SAR (Tarefa 1)", 1)
add_heading(doc, "2.1 Top 30 suspeitos", 2)
add_body(doc, (
    "Clientes foram ranqueados pela soma ponderada de regras acionadas × severidade, "
    "resultando em um score de risco de 0 a 22 (máximo). O Top 30 foi selecionado para análise aprofundada."
))

add_heading(doc, "Top 5 clientes por score (fonte: outputs/ranking_risco.csv)", 3)
add_table(
    doc,
    headers=["Rank", "Cliente", "Score", "Core", "Tipologias principais", "Volume (R$)"],
    rows=[
        ["1", "**C101208**", "22", "★",  "Sanções · Geo-salto · Cash-in→out · E-com sem 3DS · Renda incompat. · MCC risco", "150.178"],
        ["2", "C101919",     "22", "—",  "Renda incompat. · Geo-salto · Cash-in→out · Burst · E-com sem 3DS · Cross-border · PEP · MCC risco", "152.430"],
        ["3", "C100184",     "21", "—",  "Renda incompat. · Geo-salto · Anomalia IP · Cash-in→out · E-com sem 3DS · Cross-border · MCC risco", "74.888"],
        ["4", "C102093",     "21", "★",  "Renda incompat. · Geo-salto · Anomalia IP · Cash-in→out · E-com sem 3DS · Sanções · PEP · MCC risco", "192.074"],
        ["5", "C100517",     "20", "★",  "Renda incompat. · Geo-salto · Cash-in→out · Burst · E-com sem 3DS · Sanções · MCC risco", "129.922"],
    ],
    col_widths_cm=[0.8, 1.8, 1.2, 0.8, 8.4, 2.2],
)
add_body(doc, "★ = is_core_label=1 (acionou ≥2 regras-core do conjunto R03/R07/R09/R12/R15).")

doc.add_paragraph()
add_heading(doc, "Tipologias presentes no Top 30 (11 distintas)", 3)
add_table(
    doc,
    headers=["Tipologia", "Frequência no Top 30", "Exemplo"],
    rows=[
        ["Renda Incompatível (>5x renda anual)",     "predominante", "C102093 (24,7x), C101208 (11,5x)"],
        ["Geo-Salto (velocidade impossível)",        "predominante", "C101208 (1.092 km/h BR->RU em 13h)"],
        ["E-com sem 3DS (alto valor)",               "predominante", "C101208, C100517"],
        ["Cash-In → Cash-Out (PIX, 24h)",            "predominante", "C101208, C100517"],
        ["E-com sem 3DS / Cross-border (ECI)",       "alta",         "C101208, C100184"],
        ["MCC Alto Risco (6011, 7995, 6051, 4789)",  "alta",         "todo o top-5"],
        ["Sanções (R15)",                             "casos críticos", "C101208, C102093, C100517"],
        ["Anomalia IP/Device",                        "seletiva",     "C102093, C100184"],
        ["País de Alto Risco (R14)",                  "seletiva",     "C100184, C101919"],
        ["PEP",                                       "seletiva",     "C101919, C102093"],
        ["Burst/Velocidade",                          "seletiva",     "C100517, C101919"],
    ],
    col_widths_cm=[5.5, 3.5, 6.2],
)

doc.add_paragraph()
add_heading(doc, "2.2 SAR — Cliente C101208 (Caso Principal)", 2)
add_highlight_box(doc, "CRITICO — Score 22/22 — 7 Tipologias Simultâneas — Ação regulatória recomendada", "FFE6E6")

add_body(doc, (
    "Entre 2025-07-03 e 2025-09-29, o cliente C101208 (Chef, renda declarada R$ 13.047/ano, "
    "KYC Tier L1) realizou 29 transações totalizando R$ 150.178 — 11,5x sua renda anual — "
    "por PIX, Cartão e Wire em 8 países (BR, PT, BY, SY, GB, ES, RU, YE)."
))

add_heading(doc, "7 tipologias simultâneas:", 3)
typologies = [
    "Renda incompatível — volume/renda = 11,5x (R04, severidade 2)",
    "Geo-salto — 1.092 km/h BR->RU em 13h (fisicamente impossível)",
    "Cash-in -> Cash-out PIX — layering em 24h",
    "E-commerce sem 3DS — 5 transações card-not-present sem autenticação forte",
    "Cross-border + ECI não-autenticado — 9 transações internacionais",
    "Sanctions hit — TX TNHZDN7D6LYK6 com receptor em Síria (R15, severidade 4)",
    "MCC de alto risco — 6011 (ATM), 7995 (jogos/apostas), 4789 (transporte)",
]
for t in typologies:
    add_numbered(doc, t)

doc.add_paragraph()
add_heading(doc, "Ações recomendadas:", 3)
add_table(
    doc,
    headers=["Prazo", "Ação"],
    rows=[
        ["D+0", "Bloqueio preventivo da conta C101208"],
        ["D+1", "Notificação formal ao Compliance Officer"],
        ["D+3", "Avaliação de reporte ao COAF via SISCOAF (prazo legal: até 24h após decisão)"],
    ],
    col_widths_cm=[2.0, 14.2],
)

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 3 — SISTEMA DE ALERTAS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Sistema de Alertas (Tarefa 2)", 1)
add_heading(doc, "3.1 Visão geral do motor de regras", 2)
add_table(
    doc,
    headers=["Métrica", "Valor"],
    rows=[
        ["Regras implementadas",     "22"],
        ["Regras que dispararam",    "17 (5 sem ocorrências: R01, R07, R08, R11, R21)"],
        ["Total de alertas",          "10.576"],
        ["Entidades únicas no ranking","4.156 (3.310 clientes + 846 merchants)"],
        ["Weak label positivos",      "108 clientes (3,3% da carteira)"],
    ],
    col_widths_cm=[7.0, 9.2],
)

doc.add_paragraph()
add_heading(doc, "3.2 Catálogo de regras (seleção)", 2)
add_table(
    doc,
    headers=["Regra", "Tipologia", "Sev.", "Alertas", "Lógica resumida"],
    rows=[
        ["R03_structuring ★",   "Structuring",     "3",   "4",     ">=3 PIX em [9k–10k] em 7 dias"],
        ["R04_income_mismatch", "Renda x Valor",   "2",   "1.016", "Valor > 15x renda mensal"],
        ["R05_geojump",         "Geo-Salto",       "3",   "2.119", "Velocidade > 900 km/h entre tx"],
        ["R07_device_ring ★",   "Device Ring",     "3",   "0",     ">=6 clientes num device/dia"],
        ["R09_self_merchant ★", "Self-Merchant",   "4",   "2",     "Cliente paga merchant próprio"],
        ["R10_cash_in_out",     "Cash-in/out",     "3",   "802",   "Saída >= 80% entrada PIX em 24h"],
        ["R12_ecom_no_3ds ★",   "E-com sem 3DS",   "3",   "572",   "Card-NP, 3DS=No, valor alto"],
        ["R15_sanctions ★",     "Sanções",         "3–4", "484",   "sanctions_screening_hit / país sancionado"],
        ["R16_pep",             "PEP",             "3",   "52",    "pep_flag=True em tx"],
        ["R17_high_risk_mcc",   "MCC Risco",       "2",   "3.263", "MCC ∈ {6011,7995,6051,4789}"],
        ["R18_chargeback",      "Chargeback",      "2",   "846",   "Merchant com taxa CB elevada"],
    ],
    col_widths_cm=[3.5, 3.0, 1.0, 1.8, 6.9],
)
add_body(doc, "★ = regra-core (alta confiança / baixo falso-positivo) usada como rótulo fraco do ML.")

doc.add_paragraph()
add_heading(doc, "3.3 Triangulação reduz falsos positivos", 2)
add_highlight_box(
    doc,
    "R17 sozinha (MCC risco) = alerta baixa prioridade.  "
    "R17 + R05 + R15 + R04 = score 22 — nível CRITICO.",
    "E8F0FE",
)

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 4 — MODELO DE MACHINE LEARNING
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Modelo de Machine Learning (Tarefa 3)", 1)
add_heading(doc, "4.1 Arquitetura", 2)
add_body(doc, "Ensemble XGBoost (supervisionado fraco) + Isolation Forest (não-supervisionado):")
add_highlight_box(doc, "Score Final = 0,70 x XGBoost_proba + 0,30 x IF_score_normalizado", "F0F0F0")
add_body(doc, (
    "Rótulo fraco (weak label): cliente marcado positivo se acionar >= 2 regras-core de alta "
    "confiança: R03_structuring, R07_device_ring, R09_self_merchant, R12_ecom_no_3ds, "
    "R15_sanctions (definição em src/rules_engine.py, CORE_RULES, linha 82). Dataset: 3.310 "
    "clientes · 108 positivos (3,3%) · 44 features · split temporal 80/20 por first_tx_date."
))

add_heading(doc, "4.2 Métricas", 2)
add_table(
    doc,
    headers=["Métrica", "Teste único", "Cross-validation (5-fold)"],
    rows=[
        ["ROC-AUC",                    "**0,979**", "0,956 ± 0,007"],
        ["PR-AUC",                     "**0,536**", "**0,313 ± 0,048** (~9,6x baseline 0,033)"],
        ["Threshold ótimo (max-F1)",   "0,437",     "—"],
        ["Precisão @ threshold",       "0,318",     "—"],
        ["Recall @ threshold",         "**0,933** (14/15 suspeitos)", "—"],
        ["F1 @ threshold",             "0,475",     "—"],
    ],
    col_widths_cm=[5.6, 5.4, 5.4],
)

doc.add_paragraph()
add_heading(doc, "4.3 Explicabilidade — SHAP (C101208, percentil 100%)", 2)
add_table(
    doc,
    headers=["Feature", "SHAP", "Direção"],
    rows=[
        ["n_high_risk_geo",  "+0,744", "aumenta risco"],
        ["pct_high_risk_geo","+0,378", "aumenta risco"],
        ["n_no_3ds",         "+0,074", "aumenta risco"],
        ["pct_pix",          "+0,026", "aumenta risco"],
    ],
    col_widths_cm=[6.0, 2.5, 7.7],
)

doc.add_paragraph()
add_heading(doc, "4.4 Distribuição de risco da carteira", 2)
add_table(
    doc,
    headers=["Tier", "Score", "Clientes", "%"],
    rows=[
        ["Baixo",  "0–0,30", "2.824", "85,3%"],
        ["Médio",  "0,30–0,60", "237", "7,2%"],
        ["Alto",   "0,60–0,80", "249", "7,5%"],
    ],
    col_widths_cm=[3.5, 3.0, 3.0, 2.5],
)

doc.add_paragraph()
add_heading(doc, "4.5 Limitações", 2)
for lim in [
    "Rótulo ruidoso: weak label derivado de regras — o modelo aprende as regras, não lavagem real",
    "Viés circular: features correlacionadas com as regras que geraram o rótulo",
    "Janela curta (3 meses): não captura padrões sazonais ou layering multi-mês",
    "Score não calibrado: probabilidades não interpretáveis diretamente sem Platt/isotônica",
]:
    add_numbered(doc, lim)

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 5 — SISTEMA MULTI-AGENTE
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Sistema Multi-Agente (Tarefa 4)", 1)
add_heading(doc, "5.1 Arquitetura — 5 agentes + orquestrador", 2)
add_body(doc, (
    "5 agentes especializados implementados como papéis LLM via Anthropic API, chamados em "
    "sequência pelo orquestrador, que salva artefatos JSON por estágio e registra trilha de auditoria."
))

add_table(
    doc,
    headers=["#", "Agente", "Papel", "Temp."],
    rows=[
        ["1", "Dados",        "Valida schema, coerência por rail, enriquece geo/IP",      "0,0"],
        ["2", "Detecção",     "Avalia regras + ML, gera fila priorizada por score",       "0,0"],
        ["3", "Investigação", "Caso 360°: timeline, grafo, tipologias, evidências",        "0,2"],
        ["4", "Reporte",      "Rascunho de SAR em 7 seções (markdown + JSON estruturado)","0,4"],
        ["5", "Compliance",   "Valida SAR, checa sanções, decide approve/revise",         "0,0"],
    ],
    col_widths_cm=[0.8, 2.5, 11.2, 1.7],
)

doc.add_paragraph()
add_heading(doc, "5.2 Resultado do teste em tempo real — C101208", 2)
add_highlight_box(doc,
    "[1/5] DADOS       -> qualidade ok · países risco: SY, RU, YE, BY\n"
    "[2/5] DETECCAO    -> 7 regras · priority=39.6 · SANCAO no topo\n"
    "[3/5] INVESTIGACAO -> CRITICO · 7 tipologias · geo-jump 14.341km\n"
    "[4/5] REPORTE     -> SAR draft 7 secoes · layering confirmado\n"
    "[5/5] COMPLIANCE  -> APPROVE · report_coaf · SLA D+3",
    "EAF4EA",
)

add_heading(doc, "5.3 Visão de produto", 2)
add_body(doc, (
    "O pipeline representa um produto AML mínimo e completo: recebe um caso priorizado pelo "
    "motor de regras + score de ML e, em menos de 60 segundos, entrega uma decisão regulatória "
    "rastreável com base legal, trilha de auditoria e prazo de SLA — reduzindo de horas (analista "
    "manual) para segundos a triagem inicial de casos."
))

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 6 — CONCLUSÕES
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Conclusões", 1)
conclusoes = [
    ("Cobertura ampla:", "17 de 22 regras dispararam alertas, evidenciando boa cobertura das tipologias AML/FT na base CloudWalk."),
    ("Caso prioritário:", "C101208 combina 7 tipologias simultâneas em 3 rails com sanção confirmada — ação regulatória recomendada com alta confiança."),
    ("ML generaliza bem:", "ROC-AUC 0,979 e Recall 93,3% demonstram que padrões comportamentais + geográficos capturam suspeitos mesmo com rótulo fraco."),
    ("Rastreabilidade:", "O pipeline multi-agente entrega defensabilidade regulatória — cada decisão cita evidências, IDs de transação e base legal."),
    ("Limitações:", "Rótulo fraco (sem labels reais), janela de 3 meses e ausência de feedback de investigadores humanos são os principais gaps."),
]
for i, (titulo, texto) in enumerate(conclusoes, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run_t = p.add_run(titulo + " ")
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = NAVY
    run_b = p.add_run(texto)
    run_b.font.size = Pt(11)

doc.add_paragraph()
add_heading(doc, "Próximos passos para escala", 2)
for ps in [
    "Feedback loop: analista revisa -> rótulo refinado -> modelo atualizado",
    "Orquestrador com fila priorizada via LangGraph (estado compartilhado, re-enfileiramento)",
    "Graph Neural Network sobre rede de contrapartes (captura conexões cross-cliente)",
    "Monitoramento de data drift em produção (PSI, KS test)",
]:
    add_bullet(doc, ps)

# ═══════════════════════════════════════════════════════════════════════════
# SEÇÃO 7 — ANEXOS
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Anexos", 1)
add_table(
    doc,
    headers=["Artefato", "Caminho"],
    rows=[
        ["Scripts de execução (Dias 1–5)", "notebooks/01_eda_qualidade.py ... 05_agentes_pipeline.py"],
        ["Pipeline multi-agente",      "src/agents/pipeline.py"],
        ["Motor de regras",            "src/rules_engine.py"],
        ["Catálogo de regras (YAML)",  "config/rules.yaml"],
        ["SAR completo C101208 (analista)", "outputs/sar/SAR_C101208.md"],
        ["SAR gerado pelo agente (runtime)", "outputs/sar/C101208/sar_agente.md"],
        ["Scores ML (3.310 clientes)", "outputs/04_ml_scores.csv"],
        ["Ranking de risco",           "outputs/ranking_risco.csv"],
        ["Top 30 suspeitos",           "outputs/suspeitos_top30.csv"],
        ["Alertas completos",          "outputs/alertas.csv"],
        ["Figuras",                    "outputs/figuras/*.png"],
    ],
    col_widths_cm=[5.5, 10.7],
)

doc.add_paragraph()
p_disc = doc.add_paragraph()
p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_disc = p_disc.add_run(
    "Relatório gerado em 2026-06-21 — Sistema AML-FT CloudWalk — "
    "Revisão humana obrigatória antes de qualquer ação regulatória."
)
r_disc.font.size = Pt(9)
r_disc.italic = True
r_disc.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# ── Salvar ───────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"DOCX gerado: {OUT}")
